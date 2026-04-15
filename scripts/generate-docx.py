#!/usr/bin/env python3
"""
Generate Word document from memo content using Nationwide template.

Usage:
    python scripts/generate-docx.py content.json output.docx
    echo '{"to": "...", "subject": "...", "body": "..."}' | python scripts/generate-docx.py - output.docx
"""

import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def load_template():
    """Load the Nationwide memo template."""
    template_path = Path(__file__).parent.parent / "templates/pct/2024_Memo.dotx"
    if not template_path.exists():
        # Fall back to creating from scratch
        return Document()
    return Document(template_path)


def generate_memo(content, output_path):
    """Generate a Word memo from content JSON."""
    doc = load_template()

    # Clear existing content if template has any
    for para in doc.paragraphs:
        if para.text.strip():
            break  # Keep header content

    # Add memo header
    header_style = doc.styles['Normal']

    # TO line
    to_para = doc.add_paragraph()
    to_para.add_run("TO:\t\t").bold = True
    to_para.add_run(content.get("to", "[Recipients]"))

    # FROM line
    from_para = doc.add_paragraph()
    from_para.add_run("FROM:\t\t").bold = True
    from_para.add_run(content.get("from", "[Sender]"))

    # DATE line
    date_para = doc.add_paragraph()
    date_para.add_run("DATE:\t\t").bold = True
    date_para.add_run(content.get("date", datetime.now().strftime("%B %d, %Y")))

    # RE line
    re_para = doc.add_paragraph()
    re_para.add_run("RE:\t\t").bold = True
    re_para.add_run(content.get("subject", "[Subject]"))

    # Horizontal line (using paragraph border would be better, but this works)
    doc.add_paragraph("_" * 60)
    doc.add_paragraph()  # Spacing

    # Body content
    body = content.get("body", "")

    # Split body into sections if it contains headers
    lines = body.split('\n')
    current_para = None

    for line in lines:
        line = line.strip()

        if not line:
            # Empty line - add spacing
            doc.add_paragraph()
            current_para = None
        elif line.startswith('## '):
            # Section header
            para = doc.add_paragraph()
            run = para.add_run(line[3:])
            run.bold = True
            current_para = None
        elif line.startswith('- '):
            # Bullet point
            para = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            # Bold line
            para = doc.add_paragraph()
            run = para.add_run(line[2:-2])
            run.bold = True
            current_para = None
        else:
            # Regular paragraph
            if current_para:
                current_para.add_run(' ' + line)
            else:
                current_para = doc.add_paragraph(line)

    # Save
    doc.save(output_path)
    print(f"Generated: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: python generate-docx.py <content.json|-> <output.docx>")
        print("  Use '-' to read JSON from stdin")
        sys.exit(1)

    input_arg = sys.argv[1]
    output_path = sys.argv[2]

    # Read content
    if input_arg == "-":
        content = json.load(sys.stdin)
    else:
        with open(input_arg) as f:
            content = json.load(f)

    generate_memo(content, output_path)


if __name__ == "__main__":
    main()
